from pathlib import Path
import math

from PIL import Image, ImageDraw, ImageFont, ImageFilter, ImageOps
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "assets" / "booklet"
OUT_DIR.mkdir(parents=True, exist_ok=True)

IMG_REFERENCE = Path(r"C:\Users\emily\AppData\Local\Temp\codex-clipboard-6c6f55e3-9c0a-44c8-a44d-cb759aa1269c.png")
IMG_BOATS = Path(r"C:\Users\emily\Downloads\_cgi-bin_mmwebwx-bin_webwxgetmsgimg__&MsgID=4072578446302627295&skey=@crypt_dce1d0d9_6cfdf483d057c9b8c92ca4d0cd280cee&mmweb_appid=wx_webfilehelper.jpg")
IMG_CAD = Path(r"C:\Users\emily\Pictures\Screenshots\屏幕截图 2026-06-13 103052.png")

W, H = 1748, 2480  # A5 portrait at 300 dpi
DPI = (300, 300)

FONT_DIR = Path(r"C:\Windows\Fonts")
FONT_REG = FONT_DIR / "msyh.ttc"
FONT_BOLD = FONT_DIR / "msyhbd.ttc"
FONT_LATIN = FONT_DIR / "arialbd.ttf"
FONT_MONO = FONT_DIR / "consola.ttf"
if not FONT_REG.exists():
    FONT_REG = FONT_DIR / "simhei.ttf"
if not FONT_BOLD.exists():
    FONT_BOLD = FONT_REG
if not FONT_MONO.exists():
    FONT_MONO = FONT_LATIN


def font(path, size):
    return ImageFont.truetype(str(path), size)


F = {
    "title": font(FONT_LATIN, 142),
    "title_cn": font(FONT_BOLD, 76),
    "subtitle": font(FONT_REG, 42),
    "h": font(FONT_BOLD, 72),
    "body": font(FONT_REG, 36),
    "small": font(FONT_REG, 28),
    "tag": font(FONT_BOLD, 30),
    "mono": font(FONT_MONO, 24),
}


COLORS = {
    "paper": (246, 241, 231),
    "ink": (32, 42, 45),
    "muted": (94, 105, 104),
    "blue": (33, 112, 162),
    "green": (30, 146, 102),
    "orange": (222, 86, 48),
    "cream": (255, 249, 236),
    "white": (255, 255, 255),
    "dark": (20, 27, 30),
}


def cover_fill(img, size, crop_focus=(0.5, 0.5)):
    im = Image.open(img).convert("RGB")
    sw, sh = im.size
    tw, th = size
    scale = max(tw / sw, th / sh)
    nw, nh = int(sw * scale), int(sh * scale)
    im = im.resize((nw, nh), Image.LANCZOS)
    fx, fy = crop_focus
    x = int((nw - tw) * fx)
    y = int((nh - th) * fy)
    return im.crop((x, y, x + tw, y + th))


def contain(img, box, bg=(255, 255, 255), radius=28):
    im = Image.open(img).convert("RGB")
    bw, bh = box
    scale = min(bw / im.width, bh / im.height)
    nw, nh = int(im.width * scale), int(im.height * scale)
    im = im.resize((nw, nh), Image.LANCZOS)
    canvas = Image.new("RGB", box, bg)
    canvas.paste(im, ((bw - nw) // 2, (bh - nh) // 2))
    return rounded(canvas, radius)


def rounded(im, radius=30):
    mask = Image.new("L", im.size, 0)
    d = ImageDraw.Draw(mask)
    d.rounded_rectangle((0, 0, im.width, im.height), radius=radius, fill=255)
    out = Image.new("RGBA", im.size, (0, 0, 0, 0))
    out.paste(im.convert("RGBA"), (0, 0), mask)
    return out


def add_shadow(page, card, xy, blur=22, offset=(0, 12), alpha=70):
    x, y = xy
    shadow = Image.new("RGBA", card.size, (0, 0, 0, alpha))
    shadow = shadow.filter(ImageFilter.GaussianBlur(blur))
    page.alpha_composite(shadow, (x + offset[0], y + offset[1]))
    page.alpha_composite(card, xy)


def draw_text(draw, xy, text, fnt, fill, anchor=None, spacing=8, align="left"):
    draw.multiline_text(xy, text, font=fnt, fill=fill, anchor=anchor, spacing=spacing, align=align)


def wrap_text(text, fnt, max_width, draw):
    lines = []
    for raw in text.split("\n"):
        line = ""
        for ch in raw:
            test = line + ch
            if draw.textbbox((0, 0), test, font=fnt)[2] <= max_width:
                line = test
            else:
                if line:
                    lines.append(line)
                line = ch
        if line:
            lines.append(line)
    return "\n".join(lines)


def pill(draw, xy, text, fill, outline=None):
    x, y = xy
    bbox = draw.textbbox((0, 0), text, font=F["tag"])
    tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
    pad_x, pad_y = 28, 14
    rect = (x, y, x + tw + pad_x * 2, y + th + pad_y * 2)
    draw.rounded_rectangle(rect, radius=32, fill=fill, outline=outline, width=2 if outline else 0)
    draw.text((x + pad_x, y + pad_y - 2), text, font=F["tag"], fill=COLORS["white"])
    return rect[2] + 18


def page_base():
    return Image.new("RGBA", (W, H), COLORS["paper"] + (255,))


def add_page_num(draw, n):
    draw.text((W - 120, H - 90), f"{n:02d}", font=F["mono"], fill=(130, 134, 128), anchor="mm")
    draw.line((150, H - 90, W - 180, H - 90), fill=(216, 209, 194), width=3)


def page_cover():
    bg = cover_fill(IMG_BOATS, (W, H), (0.48, 0.54)).convert("RGBA")
    bg = ImageEnhanceSafe(bg, 0.82)
    overlay = Image.new("RGBA", (W, H), (12, 19, 22, 65))
    page = Image.alpha_composite(bg, overlay)
    draw = ImageDraw.Draw(page)

    # Pale title band keeps the name legible while the real prototype stays visible.
    draw.rounded_rectangle((110, 160, W - 110, 620), radius=42, fill=(246, 241, 231, 235))
    draw.text((150, 225), "Sailing", font=F["title"], fill=COLORS["ink"])
    draw.text((150, 365), "Tactic", font=F["title"], fill=COLORS["orange"])
    draw.text((155, 530), "小型智能帆船 / 原型小册", font=F["subtitle"], fill=COLORS["muted"])

    draw.rounded_rectangle((122, H - 420, W - 122, H - 135), radius=34, fill=(20, 27, 30, 196))
    copy = "从 3D 建模到真实船体\n把结构、控制与水上测试装进一条小船。"
    draw_text(draw, (170, H - 360), copy, F["body"], COLORS["cream"], spacing=14)
    pill(draw, (170, H - 205), "Prototype", COLORS["blue"])
    pill(draw, (470, H - 205), "3D Print", COLORS["green"])
    pill(draw, (770, H - 205), "Boat Control", COLORS["orange"])
    return page


def ImageEnhanceSafe(im, factor):
    return Image.blend(Image.new("RGBA", im.size, (0, 0, 0, 255)), im, factor)


def page_boats():
    page = page_base()
    draw = ImageDraw.Draw(page)
    draw.text((115, 115), "实物原型", font=F["h"], fill=COLORS["ink"])
    draw.text((118, 205), "两艘 3D 打印小船，安装舵机与控制模块，准备进入水面调试。", font=F["small"], fill=COLORS["muted"])

    photo = cover_fill(IMG_BOATS, (1518, 1320), (0.48, 0.48))
    card = rounded(photo, 34)
    add_shadow(page, card, (115, 320), blur=28, offset=(0, 16), alpha=58)

    y = 1740
    x = 118
    for text, color in [("绿色船体", COLORS["green"]), ("红色船体", COLORS["orange"]), ("模块连接", COLORS["blue"])]:
        x = pill(draw, (x, y), text, color)

    intro = "Sailing Tactic 是一个面向学习和实验的小型船控项目。我们用可打印船体、舵机和控制板，快速验证航向控制与结构设计。"
    wrapped = wrap_text(intro, F["body"], W - 240, draw)
    draw_text(draw, (118, 1875), wrapped, F["body"], COLORS["ink"], spacing=12)
    add_page_num(draw, 2)
    return page


def page_cad():
    page = page_base()
    draw = ImageDraw.Draw(page)
    draw.text((115, 115), "结构设计", font=F["h"], fill=COLORS["ink"])
    draw.text((118, 205), "在 Onshape 中完成船体、舵面、舵机支架和内部空间的组合设计。", font=F["small"], fill=COLORS["muted"])

    # Crop out most browser chrome so the CAD model becomes the hero.
    cad = Image.open(IMG_CAD).convert("RGB")
    cad = cad.crop((650, 170, 1750, 965))
    cad_card = ImageOps.contain(cad, (1518, 1180), Image.LANCZOS)
    canvas = Image.new("RGB", (1518, 1180), (255, 255, 255))
    canvas.paste(cad_card, ((1518 - cad_card.width) // 2, (1180 - cad_card.height) // 2))
    add_shadow(page, rounded(canvas, 34), (115, 340), blur=28, offset=(0, 16), alpha=48)

    notes = [
        ("01", "船体", "轻量化外壳，预留电子模块空间。"),
        ("02", "舵面", "舵机驱动方向控制，便于水面测试。"),
        ("03", "可迭代", "模型可快速修改，再打印验证。"),
    ]
    x0, y0 = 115, 1620
    gap = 28
    card_w = (W - 230 - gap * 2) // 3
    for i, (num, title, body) in enumerate(notes):
        x = x0 + i * (card_w + gap)
        draw.rounded_rectangle((x, y0, x + card_w, y0 + 430), radius=30, fill=(255, 249, 236), outline=(219, 208, 188), width=3)
        draw.text((x + 36, y0 + 38), num, font=F["mono"], fill=COLORS["orange"])
        draw.text((x + 36, y0 + 105), title, font=F["tag"], fill=COLORS["ink"])
        draw_text(draw, (x + 36, y0 + 170), wrap_text(body, F["small"], card_w - 72, draw), F["small"], COLORS["muted"], spacing=8)

    add_page_num(draw, 3)
    return page


def page_reference():
    page = page_base()
    draw = ImageDraw.Draw(page)
    draw.text((115, 115), "灵感与方向", font=F["h"], fill=COLORS["ink"])
    draw.text((118, 205), "参考开源硬件海报的表达方式，让项目看起来更像一个可以分享的作品。", font=F["small"], fill=COLORS["muted"])

    ref = contain(IMG_REFERENCE, (1518, 1320), bg=(228, 202, 183), radius=34)
    add_shadow(page, ref, (115, 325), blur=26, offset=(0, 16), alpha=48)

    draw.rounded_rectangle((115, 1750, W - 115, 2185), radius=34, fill=(32, 42, 45), outline=None)
    draw.text((165, 1810), "Sailing Tactic", font=font(FONT_LATIN, 72), fill=COLORS["cream"])
    lines = [
        "A5 小册子版本",
        "图片主导，少量说明",
        "展示从想法、建模到实物原型的完整过程",
    ]
    y = 1915
    for line in lines:
        draw.ellipse((170, y + 12, 188, y + 30), fill=COLORS["orange"])
        draw.text((210, y), line, font=F["body"], fill=COLORS["cream"])
        y += 70
    add_page_num(draw, 4)
    return page


def save_pages():
    pages = [page_cover(), page_boats(), page_cad(), page_reference()]
    paths = []
    for idx, page in enumerate(pages, 1):
        path = OUT_DIR / f"sailing_tactic_page_{idx}.png"
        page.convert("RGB").save(path, dpi=DPI, quality=95)
        paths.append(path)
    return paths


def make_docx(page_paths):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(5.83)
    section.page_height = Inches(8.27)
    section.top_margin = Inches(0)
    section.bottom_margin = Inches(0)
    section.left_margin = Inches(0)
    section.right_margin = Inches(0)
    section.header_distance = Inches(0)
    section.footer_distance = Inches(0)

    for i, path in enumerate(page_paths):
        if i:
            doc.add_page_break()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = 0
        p.paragraph_format.space_after = 0
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(5.83), height=Inches(8.27))

    out = OUT_DIR / "Sailing_Tactic_A5_Booklet.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    page_paths = save_pages()
    docx_path = make_docx(page_paths)
    print(docx_path)
