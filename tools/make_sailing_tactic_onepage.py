from pathlib import Path
import math

import numpy as np
from PIL import Image, ImageDraw, ImageFont, ImageFilter, ImageOps
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "assets" / "booklet"
OUT_DIR.mkdir(parents=True, exist_ok=True)

IMG_BOATS = Path(r"C:\Users\emily\Downloads\_cgi-bin_mmwebwx-bin_webwxgetmsgimg__&MsgID=4072578446302627295&skey=@crypt_dce1d0d9_6cfdf483d057c9b8c92ca4d0cd280cee&mmweb_appid=wx_webfilehelper.jpg")
IMG_CAD = Path(r"C:\Users\emily\Pictures\Screenshots\屏幕截图 2026-06-13 103052.png")
IMG_TOP = ROOT / "assets" / "printed-boat-top.png"

W, H = 1748, 2480  # A5 portrait at 300 dpi
DPI = (300, 300)

FONT_DIR = Path(r"C:\Windows\Fonts")
FONT_LATIN = FONT_DIR / "arialbd.ttf"
FONT_LATIN_REG = FONT_DIR / "arial.ttf"
FONT_CN = FONT_DIR / "msyh.ttc"
FONT_CN_BOLD = FONT_DIR / "msyhbd.ttc"
FONT_MONO = FONT_DIR / "consola.ttf"
if not FONT_CN.exists():
    FONT_CN = FONT_DIR / "simhei.ttf"
if not FONT_CN_BOLD.exists():
    FONT_CN_BOLD = FONT_CN
if not FONT_MONO.exists():
    FONT_MONO = FONT_LATIN_REG


def ft(path, size):
    return ImageFont.truetype(str(path), size)


F = {
    "mega": ft(FONT_LATIN, 132),
    "mega2": ft(FONT_LATIN, 128),
    "cn_title": ft(FONT_CN_BOLD, 42),
    "section": ft(FONT_LATIN, 38),
    "section_cn": ft(FONT_CN_BOLD, 35),
    "body": ft(FONT_CN, 25),
    "small": ft(FONT_CN, 21),
    "tiny": ft(FONT_CN, 18),
    "mono": ft(FONT_MONO, 19),
    "mono_big": ft(FONT_MONO, 26),
}

INK = (238, 232, 199)
PAPER = (35, 35, 34)
ORANGE = (255, 116, 32)
GREEN = (84, 211, 107)
BLUE = (0, 173, 238)
PINK = (244, 121, 160)
MUTED = (170, 166, 143)
CARD = (28, 29, 29)
BLACK = (16, 18, 18)


def rounded_mask(size, radius):
    mask = Image.new("L", size, 0)
    d = ImageDraw.Draw(mask)
    d.rounded_rectangle((0, 0, size[0], size[1]), radius=radius, fill=255)
    return mask


def cover_fill(path, size, focus=(0.5, 0.5)):
    im = Image.open(path).convert("RGB")
    scale = max(size[0] / im.width, size[1] / im.height)
    nw, nh = int(im.width * scale), int(im.height * scale)
    im = im.resize((nw, nh), Image.LANCZOS)
    x = int((nw - size[0]) * focus[0])
    y = int((nh - size[1]) * focus[1])
    return im.crop((x, y, x + size[0], y + size[1]))


def contain(path, size, bg=(255, 255, 255)):
    im = Image.open(path).convert("RGB")
    im.thumbnail(size, Image.LANCZOS)
    out = Image.new("RGB", size, bg)
    out.paste(im, ((size[0] - im.width) // 2, (size[1] - im.height) // 2))
    return out


def cutout_white_background(path):
    im = Image.open(path).convert("RGBA")
    # Crop tightly around the boat first so the source watermark/background never
    # becomes part of the foreground mask.
    w, h = im.size
    im = im.crop((int(w * 0.08), int(h * 0.24), int(w * 0.89), int(h * 0.77)))
    arr = np.array(im).astype(np.float32)
    rgb = arr[:, :, :3]
    near_white = (rgb[:, :, 0] > 232) & (rgb[:, :, 1] > 232) & (rgb[:, :, 2] > 232)
    visited = np.zeros(near_white.shape, dtype=bool)
    stack = []
    hh, ww = near_white.shape
    for x in range(ww):
        stack.append((0, x))
        stack.append((hh - 1, x))
    for y in range(hh):
        stack.append((y, 0))
        stack.append((y, ww - 1))
    while stack:
        y, x = stack.pop()
        if y < 0 or y >= hh or x < 0 or x >= ww or visited[y, x] or not near_white[y, x]:
            continue
        visited[y, x] = True
        stack.extend(((y - 1, x), (y + 1, x), (y, x - 1), (y, x + 1)))
    alpha_arr = np.where(visited, 0, 255).astype(np.uint8)
    alpha = Image.fromarray(alpha_arr, "L").filter(ImageFilter.GaussianBlur(1.1))
    im.putalpha(alpha)
    bbox = alpha.getbbox()
    if bbox:
        im = im.crop(bbox)
    return im


def shadow_paste(base, im, xy, blur=18, alpha=120):
    x, y = xy
    shadow = Image.new("RGBA", im.size, (0, 0, 0, 0))
    shadow_alpha = im.getchannel("A").filter(ImageFilter.GaussianBlur(blur)).point(lambda v: int(v * alpha / 255))
    shadow.putalpha(shadow_alpha)
    base.alpha_composite(shadow, (x + 18, y + 20))
    base.alpha_composite(im, xy)


def wrap(draw, text, font, width):
    lines = []
    for para in text.split("\n"):
        line = ""
        for ch in para:
            test = line + ch
            if draw.textbbox((0, 0), test, font=font)[2] <= width:
                line = test
            else:
                if line:
                    lines.append(line)
                line = ch
        if line:
            lines.append(line)
    return "\n".join(lines)


def draw_block(draw, xy, wh, title, lines, accent=ORANGE):
    x, y = xy
    w, h = wh
    draw.rounded_rectangle((x, y, x + w, y + h), radius=10, outline=accent, width=4, fill=(32, 32, 31))
    title_w = draw.textbbox((0, 0), title, font=F["section"])[2]
    draw.rectangle((x + 24, y - 8, x + 42 + title_w, y + 17), fill=PAPER)
    draw.text((x + 30, y - 17), title, font=F["section"], fill=accent)
    yy = y + 42
    for line in lines:
        if not line:
            yy += 17
            continue
        wrapped = wrap(draw, line, F["small"], w - 56)
        draw.multiline_text((x + 28, yy), wrapped, font=F["small"], fill=INK, spacing=7)
        line_count = max(1, wrapped.count("\n") + 1)
        yy += 34 * line_count + 8


def make_poster():
    page = Image.new("RGBA", (W, H), PAPER + (255,))
    draw = ImageDraw.Draw(page)

    # Subtle technical texture inspired by the reference poster.
    for yy in range(0, H, 34):
        for xx in range(0, W, 34):
            if (xx // 34 + yy // 34) % 2 == 0:
                draw.ellipse((xx + 7, yy + 7, xx + 11, yy + 11), fill=(48, 48, 46))
    for i in range(36):
        y = 260 + i * 38
        draw.line((100, y, 540, y - 165), fill=(49, 50, 48), width=9)
        draw.line((1210, y + 700, 1640, y + 535), fill=(46, 47, 46), width=7)

    # Title.
    draw.text((118, 70), "SAILING", font=F["mega"], fill=INK)
    draw.text((118, 202), "TACTIC", font=F["mega2"], fill=INK)
    draw.line((125, 355, W - 125, 355), fill=ORANGE, width=10)
    draw.ellipse((96, 340, 124, 368), fill=ORANGE)
    draw.ellipse((W - 124, 340, W - 96, 368), fill=ORANGE)
    draw.text((126, 385), "3D PRINTED SAILING CONTROLLER / BLE KEYBOARD FOR TACTICAL SAILING",
              font=F["mono_big"], fill=INK)
    draw.text((128, 428), "实体小船旋钮控制屏幕里的帆船，支持双船对战、校准和 LED 状态提示",
              font=F["cn_title"], fill=GREEN)

    # Large cutout window: center-right balances the title and gives the object room.
    win = (565, 545, 1625, 1295)
    draw.rounded_rectangle(win, radius=22, fill=BLACK, outline=ORANGE, width=7)
    for k in range(0, 740, 54):
        draw.line((win[0] + 18, win[1] + k, win[2] - 18, win[1] + k - 260),
                  fill=(55, 58, 57), width=4)
    cad_crop = Image.open(IMG_CAD).convert("RGB").crop((650, 170, 1790, 1025))
    cad_crop = ImageOps.contain(cad_crop, (920, 560), Image.LANCZOS)
    cad_layer = Image.new("RGBA", (920, 560), (0, 0, 0, 0))
    cad_layer.alpha_composite(cad_crop.convert("RGBA").point(lambda p: p))
    cad_layer = cad_layer.filter(ImageFilter.GaussianBlur(0.3))
    page.alpha_composite(cad_layer.copy().putalpha(70) if False else Image.blend(Image.new("RGBA", cad_layer.size, (0, 0, 0, 0)), cad_layer, 0.42),
                         (690, 640))
    boat = cutout_white_background(IMG_TOP)
    boat = boat.resize((980, int(980 * boat.height / boat.width)), Image.LANCZOS)
    shadow_paste(page, boat, (620, 760), blur=20, alpha=140)
    draw.text((600, 588), "LARGE CUTOUT WINDOW", font=F["mono_big"], fill=BLUE)
    draw.text((600, 625), "top-view prototype + CAD layer", font=F["mono"], fill=MUTED)

    # Left intro panel.
    intro = [
        "SAILING TACTIC 是一个用实体小船操控 Tactical Sailing 的",
        "桌面控制器。M5Stack Chain Angle 读取舵角，Atom Basic",
        "通过 BLE Keyboard 把动作转换成电脑按键。",
        "",
        "旋转船上的舵柄，屏幕里的船就跟着转向。"
    ]
    draw_block(draw, (118, 650), (385, 480), "WHAT", intro, ORANGE)

    # Photo strip with actual two-boat build.
    photo = cover_fill(IMG_BOATS, (620, 360), (0.52, 0.48)).convert("RGBA")
    mask = rounded_mask(photo.size, 16)
    photo.putalpha(mask)
    page.alpha_composite(photo, (118, 1195))
    draw.rounded_rectangle((118, 1195, 738, 1555), radius=16, outline=GREEN, width=5)
    draw.text((140, 1572), "REAL BUILD: dual 3D printed boat prototype",
              font=F["mono"], fill=GREEN)

    # Feature blocks.
    draw_block(draw, (118, 1670), (470, 585), "FEATURES", [
        "• 3D 打印船体和舵柄",
        "• M5Stack Atom Basic 主控",
        "• Chain Angle 读取 12-bit 舵角",
        "• HY2.0-4P 链式连接，少接线",
        "• 蓝牙键盘模式控制电脑游戏",
        "• 最多扫描 4 个角度传感器"
    ], GREEN)

    draw_block(draw, (630, 1445), (438, 405), "CONTROLS", [
        "短按: 开启 / 关闭游戏模式",
        "长按 5 秒: 进入角度校准",
        "校准中双击: 保存中心点",
        "长按 10 秒: 清除蓝牙配对",
        "A1: Left / Right",
        "A2: X / V"
    ], BLUE)

    draw_block(draw, (1110, 1445), (515, 405), "LED STATUS", [
        "红色: 游戏模式正在发送按键",
        "黄色闪烁: 等待蓝牙连接",
        "紫色闪烁: 校准模式",
        "蓝色 / 青绿色: 蓝牙连接状态",
        "传感器灯越靠近中心越亮"
    ], PINK)

    draw_block(draw, (630, 1940), (995, 315), "WHY IT EXISTS", [
        "给帆船俱乐部、招生活动和雨天训练用的教学小道具。",
        "路过的小朋友可以坐下来，用真实小船遥控电脑里的船，",
        "一边玩 Tactical Sailing，一边理解推舵、拉舵和航向变化。"
    ], ORANGE)

    # Small hardware note and cost.
    draw.rounded_rectangle((118, 2285, 1625, 2375), radius=10, fill=(26, 26, 25), outline=(68, 66, 58), width=3)
    draw.text((145, 2312), "BOM ≈ RMB 157  /  ESP32 + NeoPixel + M5Chain + BLE Keyboard  /  CAD: Onshape shared design",
              font=F["mono"], fill=INK)

    # Reference-style side labels and doodle marks.
    draw.text((1635, 610), "prototype_v1.0", font=F["mono"], fill=MUTED, anchor="mm")
    for p1, p2, c in [
        ((1490, 510), (1570, 455), GREEN),
        ((510, 595), (570, 540), BLUE),
        ((750, 1345), (830, 1310), ORANGE),
        ((430, 1618), (508, 1575), PINK),
    ]:
        draw.line((*p1, *p2), fill=c, width=5)
        draw.ellipse((p2[0] - 8, p2[1] - 8, p2[0] + 8, p2[1] + 8), outline=c, width=4)

    out_png = OUT_DIR / "Sailing_Tactic_A5_OnePage_Poster.png"
    page.convert("RGB").save(out_png, dpi=DPI, quality=95)
    return out_png


def make_docx(poster_path):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(5.83)
    section.page_height = Inches(8.27)
    section.top_margin = Inches(0)
    section.bottom_margin = Inches(0)
    section.left_margin = Inches(0)
    section.right_margin = Inches(0)
    section.header_distance = Inches(0)
    section.footer_distance = Inches(0)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = 0
    p.paragraph_format.space_after = 0
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(poster_path), width=Inches(5.83), height=Inches(8.27))
    out_docx = OUT_DIR / "Sailing_Tactic_A5_OnePage_Poster.docx"
    doc.save(out_docx)
    return out_docx


if __name__ == "__main__":
    poster = make_poster()
    docx = make_docx(poster)
    print(docx)
    print(poster)
